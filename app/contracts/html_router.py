from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import Response
from sqlalchemy.orm import Session
from typing import Optional
import os
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

from app.dependencies.database.database import get_db
from app.auth.dependencies.get_current_user import get_current_user
from app.models.user_model import User
from app.models.history_model import RentalHistory
from app.models.car_model import Car
from app.utils.short_id import safe_sid_to_uuid

HTMLContractsRouter = APIRouter(prefix="/contracts/html", tags=["HTML Contracts"])

CONTRACT_FILES = {
    "rental_main_contract": "rental_main_contract.html",
    "appendix_7_1": "acceptance_certificate.html",
    "appendix_7_2": "return_certificate.html",
    "main_contract": "main_contract.html",
    "user_agreement": "user_agreement.html",
    "consent_to_data_processing": "consent_to_the_processing_of_personaldata.html",
    "main_contract_for_guarantee": "main_contract_for_guarantee.html",
}

UPLOADS_DOCS_PATH = "uploads/docs"


def format_data(value: Optional[str], fallback: str = "не указано") -> str:
    """Форматирует данные с fallback значением"""
    if value is None or value == "":
        return fallback
    return str(value)


def format_car_data(value: Optional[str], fallback: str = "данные не указаны") -> str:
    """Форматирует данные автомобиля"""
    return format_data(value, fallback)


def format_year(value: Optional[str], fallback: str = "выпуска: не указан") -> str:
    """Форматирует год выпуска"""
    return format_data(value, fallback)


def format_vin(value: Optional[str], fallback: str = "VIN не указан") -> str:
    """Форматирует VIN"""
    return format_data(value, fallback)


def format_color(value: Optional[str], fallback: str = "Цвет: не указан") -> str:
    """Форматирует цвет"""
    return format_data(value, fallback)


def format_car_uuid(value: Optional[str], fallback: str = "UUID не указан") -> str:
    """Форматирует UUID автомобиля"""
    return format_data(value, fallback)


def get_current_date() -> str:
    """Возвращает текущую дату в формате DD.MM.YYYY"""
    return datetime.now().strftime("%d.%m.%Y")


def translate_body_type(body_type: Optional[str]) -> str:
    """Переводит тип кузова на русский"""
    if not body_type:
        return "не указан"
    
    body_type_map = {
        "SEDAN": "Седан",
        "SUV": "Внедорожник",
        "CROSSOVER": "Кроссовер",
        "COUPE": "Купе",
        "HATCHBACK": "Хэтчбек",
        "CONVERTIBLE": "Кабриолет",
        "WAGON": "Универсал",
        "MINIBUS": "Микроавтобус",
        "ELECTRIC": "Электромобиль",
    }
    
    return body_type_map.get(body_type.upper(), body_type)


def html_to_docx(html_content: str) -> BytesIO:
    """Конвертирует HTML в Word документ (.docx)"""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    doc = Document()
    
    for script in soup(["script", "style"]):
        script.decompose()
    
    body = soup.find('body') or soup
    
    def add_element_to_doc(element, doc):
        """Рекурсивно добавляет элементы в документ"""
        if element.name is None: 
            text = str(element).strip()
            if text:
                if doc.paragraphs:
                    doc.paragraphs[-1].add_run(text)
                else:
                    p = doc.add_paragraph(text)
        elif element.name in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            text = element.get_text(strip=True)
            if text:
                p = doc.add_paragraph(text)
                if element.name.startswith('h'):
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(14 if element.name == 'h1' else 12)
        elif element.name == 'br':
            if doc.paragraphs:
                doc.paragraphs[-1].add_run().add_break()
        elif element.name == 'table':
            table = doc.add_table(rows=0, cols=0)
            for tr in element.find_all('tr'):
                row_cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                if row_cells:
                    row = table.add_row()
                    for i, cell_text in enumerate(row_cells):
                        if i < len(row.cells):
                            row.cells[i].text = cell_text
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                text = li.get_text(strip=True)
                if text:
                    p = doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
        else:
            for child in element.children:
                add_element_to_doc(child, doc)
    
    for element in body.children:
        add_element_to_doc(element, doc)
    
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer


def process_html_placeholders(
    html: str,
    full_name: Optional[str] = None,
    login: Optional[str] = None,
    client_id: Optional[str] = None,
    digital_signature: Optional[str] = None,
    rental_id: Optional[str] = None,
    car_name: Optional[str] = None,
    plate_number: Optional[str] = None,
    car_uuid: Optional[str] = None,
    car_year: Optional[str] = None,
    body_type: Optional[str] = None,
    vin: Optional[str] = None,
    color: Optional[str] = None,
) -> str:
    """Обрабатывает плейсхолдеры в HTML"""
    
    html = html.replace("${full_name}", format_car_data(full_name, "ФИО не указано"))
    html = html.replace("${login}", format_car_data(login, "логин не указан"))
    html = html.replace("${client_id}", format_car_data(client_id, "ID не указан"))
    html = html.replace("${client_uuid}", format_car_data(client_id, "ID не указан"))
    html = html.replace("${digital_signature}", format_car_data(digital_signature, "подпись не указана"))
    html = html.replace("${rental_id}", format_car_data(rental_id, "ID аренды не указан"))
    html = html.replace("${rent_uuid}", format_car_data(rental_id, "ID аренды не указан"))
    html = html.replace("${rent_id}", format_car_data(rental_id, "ID аренды не указан"))
    html = html.replace("${car_name}", format_car_data(car_name, "не указано"))
    html = html.replace("${plate_number}", format_car_data(plate_number, "не указан"))
    html = html.replace("${car_uuid}", format_car_data(car_uuid, "не указан"))
    html = html.replace("${car_year}", format_year(car_year))
    html = html.replace("${body_type}", translate_body_type(body_type))
    html = html.replace("${vin}", format_vin(vin))
    html = html.replace("${color}", format_color(color))
    html = html.replace("${date}", get_current_date())
    html = html.replace("{date}", get_current_date())
    
    html = html.replace("{___car_name_____}", format_car_data(car_name, "не указано"))
    html = html.replace("{____car_plate_number______}", format_car_data(plate_number, "не указан"))
    html = html.replace("{_______car_id________}", format_car_uuid(car_uuid, "не указан"))
    html = html.replace("{_____car_year___________}", format_year(car_year))
    html = html.replace("{______car_body_type_____________}", translate_body_type(body_type))
    html = html.replace("{______car_vin_________}", format_vin(vin))
    html = html.replace("{________car_color_________}", format_color(color))
    
    html = html.replace("{___________________}", format_car_data(body_type, "не указан"))
    html = html.replace("выпуска: {________________}", format_year(car_year))
    html = html.replace("№: {_______________}", format_vin(vin))
    html = html.replace("{__________}", format_car_data(plate_number, "не указан"))
    html = html.replace("{________}", format_car_data(car_name, "не указано"))
    html = html.replace("цвет: {________________}", format_color(color))
    html = html.replace("номер: {_______________}", format_car_uuid(car_uuid, "не указан"))
    
    if "viewport" not in html.lower():
        viewport_meta = """<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes">"""
        styles = """
        <style>
            * {
                box-sizing: border-box;
            }
            html, body {
                margin: 0;
                padding: 0;
                width: 100%;
                overflow-x: hidden;
                -webkit-overflow-scrolling: touch;
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            }
            body {
                padding: 16px;
                line-height: 1.6;
                font-size: 14px;
                background: white;
            }
            p, div, span, td, th {
                -webkit-user-select: text;
                user-select: text;
            }
            img {
                max-width: 100%;
                height: auto;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                font-size: 13px;
            }
            @media screen and (max-width: 768px) {
                body {
                    font-size: 13px;
                    padding: 12px;
                }
                table {
                    font-size: 12px;
                }
            }
        </style>"""
        
        if "<head>" in html:
            html = html.replace("<head>", f"<head>\n{viewport_meta}\n{styles}")
        else:
            html = f"<head>\n{viewport_meta}\n{styles}\n</head>\n{html}"
    
    return html


@HTMLContractsRouter.get("/rental-main-contract")
async def get_rental_main_contract(
    full_name: Optional[str] = Query(None, description="ФИО клиента"),
    login: Optional[str] = Query(None, description="Логин клиента"),
    client_id: Optional[str] = Query(None, description="ID клиента"),
    digital_signature: Optional[str] = Query(None, description="Цифровая подпись"),
    rental_id: Optional[str] = Query(None, description="ID аренды"),
    car_name: Optional[str] = Query(None, description="Марка и модель автомобиля"),
    plate_number: Optional[str] = Query(None, description="Гос. номер"),
    car_uuid: Optional[str] = Query(None, description="UUID автомобиля"),
    car_year: Optional[str] = Query(None, description="Год выпуска"),
    body_type: Optional[str] = Query(None, description="Тип кузова"),
    vin: Optional[str] = Query(None, description="VIN номер"),
    color: Optional[str] = Query(None, description="Цвет"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Получить договор присоединения для аренды с подставленными данными"""
    file_path = os.path.join(UPLOADS_DOCS_PATH, CONTRACT_FILES["rental_main_contract"])
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл контракта не найден"
        )
    
    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()
    
    if rental_id:
        try:
            rental_uuid = safe_sid_to_uuid(rental_id)
            rental = db.query(RentalHistory).filter(RentalHistory.id == rental_uuid).first()
            
            if rental:
                if rental.user_id != current_user.id and current_user.role.value not in ["ADMIN", "MECHANIC"]:
                    raise HTTPException(
                        status_code=status.HTTP_403_FORBIDDEN,
                        detail="Доступ запрещен"
                    )
                
                car = db.query(Car).filter(Car.id == rental.car_id).first()
                if car:
                    car_name = car_name or car.name
                    plate_number = plate_number or car.plate_number
                    car_uuid = car_uuid or str(car.id)
                    car_year = car_year or (str(car.year) if car.year else None)
                    body_type = body_type or (car.body_type.value if car.body_type else None)
                    vin = vin or car.vin
                    color = color or car.color
        except Exception as e:
            pass  
    
    html = process_html_placeholders(
        html,
        full_name=full_name or current_user.first_name + " " + current_user.last_name if current_user.first_name and current_user.last_name else None,
        login=login or current_user.phone_number,
        client_id=client_id or str(current_user.id),
        digital_signature=digital_signature or current_user.digital_signature,
        rental_id=rental_id,
        car_name=car_name,
        plate_number=plate_number,
        car_uuid=car_uuid,
        car_year=car_year,
        body_type=body_type,
        vin=vin,
        color=color,
    )
    
    docx_buffer = html_to_docx(html)
    
    return Response(
        content=docx_buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=contract.docx",
            "Cache-Control": "no-cache, no-store, must-revalidate"
        }
    )


@HTMLContractsRouter.get("/acceptance-certificate")
async def get_acceptance_certificate(
    full_name: Optional[str] = Query(None),
    login: Optional[str] = Query(None),
    client_id: Optional[str] = Query(None),
    digital_signature: Optional[str] = Query(None),
    rental_id: Optional[str] = Query(None),
    car_name: Optional[str] = Query(None),
    plate_number: Optional[str] = Query(None),
    car_uuid: Optional[str] = Query(None),
    car_year: Optional[str] = Query(None),
    body_type: Optional[str] = Query(None),
    vin: Optional[str] = Query(None),
    color: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Получить акт приема-передачи (Приложение 7.1) с подставленными данными"""
    file_path = os.path.join(UPLOADS_DOCS_PATH, CONTRACT_FILES["appendix_7_1"])
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл контракта не найден"
        )
    
    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()
    
    if rental_id:
        try:
            rental_uuid = safe_sid_to_uuid(rental_id)
            rental = db.query(RentalHistory).filter(RentalHistory.id == rental_uuid).first()
            
            if rental:
                if rental.user_id != current_user.id and current_user.role.value not in ["ADMIN", "MECHANIC"]:
                    raise HTTPException(
                        status_code=status.HTTP_403_FORBIDDEN,
                        detail="Доступ запрещен"
                    )
                
                car = db.query(Car).filter(Car.id == rental.car_id).first()
                if car:
                    car_name = car_name or car.name
                    plate_number = plate_number or car.plate_number
                    car_uuid = car_uuid or str(car.id)
                    car_year = car_year or (str(car.year) if car.year else None)
                    body_type = body_type or (car.body_type.value if car.body_type else None)
                    vin = vin or car.vin
                    color = color or car.color
        except Exception:
            pass
    
    html = process_html_placeholders(
        html,
        full_name=full_name or (current_user.first_name + " " + current_user.last_name if current_user.first_name and current_user.last_name else None),
        login=login or current_user.phone_number,
        client_id=client_id or str(current_user.id),
        digital_signature=digital_signature or current_user.digital_signature,
        rental_id=rental_id,
        car_name=car_name,
        plate_number=plate_number,
        car_uuid=car_uuid,
        car_year=car_year,
        body_type=body_type,
        vin=vin,
        color=color,
    )
    
    docx_buffer = html_to_docx(html)
    
    return Response(
        content=docx_buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=contract.docx",
            "Cache-Control": "no-cache, no-store, must-revalidate"
        }
    )


@HTMLContractsRouter.get("/return-certificate")
async def get_return_certificate(
    full_name: Optional[str] = Query(None),
    login: Optional[str] = Query(None),
    client_id: Optional[str] = Query(None),
    digital_signature: Optional[str] = Query(None),
    rental_id: Optional[str] = Query(None),
    car_name: Optional[str] = Query(None),
    plate_number: Optional[str] = Query(None),
    car_uuid: Optional[str] = Query(None),
    car_year: Optional[str] = Query(None),
    body_type: Optional[str] = Query(None),
    vin: Optional[str] = Query(None),
    color: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Получить акт возврата (Приложение 7.2) с подставленными данными"""
    file_path = os.path.join(UPLOADS_DOCS_PATH, CONTRACT_FILES["appendix_7_2"])
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл контракта не найден"
        )
    
    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()
    
    if rental_id:
        try:
            rental_uuid = safe_sid_to_uuid(rental_id)
            rental = db.query(RentalHistory).filter(RentalHistory.id == rental_uuid).first()
            
            if rental:
                if rental.user_id != current_user.id and current_user.role.value not in ["ADMIN", "MECHANIC"]:
                    raise HTTPException(
                        status_code=status.HTTP_403_FORBIDDEN,
                        detail="Доступ запрещен"
                    )
                
                car = db.query(Car).filter(Car.id == rental.car_id).first()
                if car:
                    car_name = car_name or car.name
                    plate_number = plate_number or car.plate_number
                    car_uuid = car_uuid or str(car.id)
                    car_year = car_year or (str(car.year) if car.year else None)
                    body_type = body_type or (car.body_type.value if car.body_type else None)
                    vin = vin or car.vin
                    color = color or car.color
        except Exception:
            pass
    
    html = process_html_placeholders(
        html,
        full_name=full_name or (current_user.first_name + " " + current_user.last_name if current_user.first_name and current_user.last_name else None),
        login=login or current_user.phone_number,
        client_id=client_id or str(current_user.id),
        digital_signature=digital_signature or current_user.digital_signature,
        rental_id=rental_id,
        car_name=car_name,
        plate_number=plate_number,
        car_uuid=car_uuid,
        car_year=car_year,
        body_type=body_type,
        vin=vin,
        color=color,
    )
    
    docx_buffer = html_to_docx(html)
    
    return Response(
        content=docx_buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=contract.docx",
            "Cache-Control": "no-cache, no-store, must-revalidate"
        }
    )


@HTMLContractsRouter.get("/main-contract")
async def get_main_contract(
    full_name: Optional[str] = Query(None),
    login: Optional[str] = Query(None),
    client_id: Optional[str] = Query(None),
    digital_signature: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
):
    """Получить основной договор с подставленными данными"""
    file_path = os.path.join(UPLOADS_DOCS_PATH, CONTRACT_FILES["main_contract"])
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл контракта не найден"
        )
    
    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()
    
    html = process_html_placeholders(
        html,
        full_name=full_name or (current_user.first_name + " " + current_user.last_name if current_user.first_name and current_user.last_name else None),
        login=login or current_user.phone_number,
        client_id=client_id or str(current_user.id),
        digital_signature=digital_signature or current_user.digital_signature,
    )
    
    docx_buffer = html_to_docx(html)
    
    return Response(
        content=docx_buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=contract.docx",
            "Cache-Control": "no-cache, no-store, must-revalidate"
        }
    )


@HTMLContractsRouter.get("/user-agreement")
async def get_user_agreement(
    full_name: Optional[str] = Query(None),
    login: Optional[str] = Query(None),
    client_id: Optional[str] = Query(None),
    digital_signature: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
):
    """Получить пользовательское соглашение с подставленными данными"""
    file_path = os.path.join(UPLOADS_DOCS_PATH, CONTRACT_FILES["user_agreement"])
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл контракта не найден"
        )
    
    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()
    
    html = process_html_placeholders(
        html,
        full_name=full_name or (current_user.first_name + " " + current_user.last_name if current_user.first_name and current_user.last_name else None),
        login=login or current_user.phone_number,
        client_id=client_id or str(current_user.id),
        digital_signature=digital_signature or current_user.digital_signature,
    )
    
    docx_buffer = html_to_docx(html)
    
    return Response(
        content=docx_buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=contract.docx",
            "Cache-Control": "no-cache, no-store, must-revalidate"
        }
    )


@HTMLContractsRouter.get("/consent-to-data-processing")
async def get_consent_to_data_processing(
    full_name: Optional[str] = Query(None),
    login: Optional[str] = Query(None),
    client_id: Optional[str] = Query(None),
    digital_signature: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
):
    """Получить согласие на обработку персональных данных с подставленными данными"""
    file_path = os.path.join(UPLOADS_DOCS_PATH, CONTRACT_FILES["consent_to_data_processing"])
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл контракта не найден"
        )
    
    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()
    
    html = process_html_placeholders(
        html,
        full_name=full_name or (current_user.first_name + " " + current_user.last_name if current_user.first_name and current_user.last_name else None),
        login=login or current_user.phone_number,
        client_id=client_id or str(current_user.id),
        digital_signature=digital_signature or current_user.digital_signature,
    )
    
    docx_buffer = html_to_docx(html)
    
    return Response(
        content=docx_buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=contract.docx",
            "Cache-Control": "no-cache, no-store, must-revalidate"
        }
    )


@HTMLContractsRouter.get("/main-contract-for-guarantee")
async def get_main_contract_for_guarantee(
    full_name: Optional[str] = Query(None),
    login: Optional[str] = Query(None),
    client_id: Optional[str] = Query(None),
    digital_signature: Optional[str] = Query(None),
    guarantor_fullname: Optional[str] = Query(None),
    guarantor_iin: Optional[str] = Query(None),
    guarantor_phone: Optional[str] = Query(None),
    guarantor_email: Optional[str] = Query(None),
    guarantor_id: Optional[str] = Query(None),
    client_fullname: Optional[str] = Query(None),
    client_iin: Optional[str] = Query(None),
    client_phone: Optional[str] = Query(None),
    client_email: Optional[str] = Query(None),
    guarantee_id: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
):
    """Получить основной договор для гаранта с подставленными данными"""
    file_path = os.path.join(UPLOADS_DOCS_PATH, CONTRACT_FILES["main_contract_for_guarantee"])
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Файл контракта не найден"
        )
    
    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()
    
    html = html.replace("{_____guarantor_fullname_______}", format_car_data(guarantor_fullname))
    html = html.replace("{____guarantor_iin________}", format_car_data(guarantor_iin))
    html = html.replace("{_____guarantor_phone_______}", format_car_data(guarantor_phone))
    html = html.replace("{____guarantor_phone____}", format_car_data(guarantor_phone))
    html = html.replace("{__ guarantor_email__}", format_car_data(guarantor_email))
    html = html.replace("{___guarantor_id____}", format_car_data(guarantor_id))
    
    html = html.replace("{_____client_fullname_______}", format_car_data(client_fullname))
    html = html.replace("{___client_iin____}", format_car_data(client_iin))
    html = html.replace("{___client_phone______}", format_car_data(client_phone))
    html = html.replace("{____client_email______}", format_car_data(client_email))
    html = html.replace("{__client_id___}", format_car_data(client_id))
    
    html = html.replace("${full_name}", format_car_data(guarantor_fullname or full_name))
    html = html.replace("${login}", format_car_data(login or guarantor_phone or current_user.phone_number))
    html = html.replace("${client_id}", format_car_data(client_id))
    html = html.replace("${digital_signature}", format_car_data(digital_signature or current_user.digital_signature))
    html = html.replace("${guarantee_id}", format_car_data(guarantee_id))
    html = html.replace("${guarantor_id}", format_car_data(guarantor_id))
    html = html.replace("${guarantor_phone}", format_car_data(guarantor_phone))
    html = html.replace("${renter}", format_car_data(client_fullname))
    html = html.replace("${date}", get_current_date())
    html = html.replace("{date}", get_current_date())
    
    if "viewport" not in html.lower():
        viewport_meta = """<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=5.0, user-scalable=yes">"""
        styles = """
        <style>
            * {
                box-sizing: border-box;
            }
            html, body {
                margin: 0;
                padding: 0;
                width: 100%;
                overflow-x: hidden;
                -webkit-overflow-scrolling: touch;
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            }
            body {
                padding: 16px;
                line-height: 1.6;
                font-size: 14px;
                background: white;
            }
            p, div, span, td, th {
                -webkit-user-select: text;
                user-select: text;
            }
            img {
                max-width: 100%;
                height: auto;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                font-size: 13px;
            }
            @media screen and (max-width: 768px) {
                body {
                    font-size: 13px;
                    padding: 12px;
                }
                table {
                    font-size: 12px;
                }
            }
        </style>"""
        
        if "<head>" in html:
            html = html.replace("<head>", f"<head>\n{viewport_meta}\n{styles}")
        else:
            html = f"<head>\n{viewport_meta}\n{styles}\n</head>\n{html}"
    
    docx_buffer = html_to_docx(html)
    
    return Response(
        content=docx_buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=contract.docx",
            "Cache-Control": "no-cache, no-store, must-revalidate"
        }
    )

